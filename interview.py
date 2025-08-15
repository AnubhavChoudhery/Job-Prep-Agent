import requests
from docx import Document
from datetime import datetime
from dotenv import load_dotenv
import os
import re
from pathlib import Path
from ibm_watsonx_ai import Credentials
from ibm_watsonx_ai.foundation_models import ModelInference
from ibm_watsonx_ai.metanames import GenTextParamsMetaNames
from location import get_location
import warnings

warnings.filterwarnings("ignore")
def init_model(WATSONX_URL, WATSONX_API_KEY, WATSONX_PROJECT_ID, WATSONX_MODEL_ID) -> ModelInference:
    creds = Credentials(url=WATSONX_URL, api_key=WATSONX_API_KEY)
    return ModelInference(
        model_id=WATSONX_MODEL_ID,
        credentials=creds,
        project_id=WATSONX_PROJECT_ID
    )

def generate_text(model: ModelInference, system_prompt: str, user_text: str, max_tokens=2000):
    """Generate text response from the model (not JSON)"""
    params = {
        GenTextParamsMetaNames.MAX_NEW_TOKENS: max_tokens,
        GenTextParamsMetaNames.TEMPERATURE: 0.2,
        GenTextParamsMetaNames.DECODING_METHOD: "greedy",
    }
    prompt = f"{system_prompt}\n\n=== INPUT START ===\n{user_text}\n=== INPUT END ===="
    
    try:
        out = model.generate_text(prompt=prompt, params=params)
        
        if isinstance(out, str):
            return out.strip()
        elif isinstance(out, dict):
            return out.get("results", [{}])[0].get("generated_text", "").strip()
        else:
            return str(out).strip()
    except Exception as e:
        print(f"[ERROR] LLM generation failed: {e}")
        return ""

def search(position, company, code, api_key, url):
    query_text = f"{position} interview experience {company}"
    params = {
        "api_key": api_key,
        "query": query_text,
        "country": code
    }
    print(f"Searching for: {query_text}")
    response = requests.get(url, params=params)
    
    if response.status_code == 200:
        return response
    else:
        raise RuntimeError(f"Scraping request failed ({response.status_code})")

def extract_text(response):
    """Extract readable text content from the API response"""
    try:
        data = response.json()
        text_content = ""
        if 'text_blocks' in data:
            for block in data['text_blocks']:
                if block.get('type') == 'paragraph':
                    text_content += block.get('snippet', '') + "\n\n"
                elif block.get('type') == 'list' and 'items' in block:
                    for item in block['items']:
                        text_content += "• " + item.get('snippet', '') + "\n"
                    text_content += "\n"
        
        return text_content.strip()
    except Exception as e:
        print(f"{e}")
        return str(response.text)

def summarize_content(model, content):
    QA_PROMPT = """You are an expert career guide tasked with summarizing interview experiences.
    Extract key points from the provided data to generate a detailed and well-structured document.
    
    Structure your output as follows:
    INTERVIEW PROCESS OVERVIEW
    • [Key points about the overall process]
    
    INTERVIEW ROUNDS AND FORMAT
    • [Details about different rounds]
    
    COMMONLY ASKED QUESTIONS AND TOPICS
    • [Types of questions and technical topics]
    
    DIFFICULTY LEVEL AND EXPECTATIONS
    • [Assessment of difficulty and what interviewers look for]
    
    PREPARATION TIPS AND RECOMMENDATIONS
    • [Specific advice for preparation]
    
    KEY TAKEAWAYS
    • [Important insights and final advice]
    Make sure to be comprehensive and specific based on the provided content.,
    diving as deep as possible into specifics at each step"""
    
    try:
        result = generate_text(model, QA_PROMPT, content, max_tokens=2000)
        return result
    except Exception as e:
        print(f"{e}")
        exit(1)

def sanitize_filename(filename):
    """Remove invalid characters from a filename."""
    return re.sub(r'[\\/:*?"<>|()]', '', filename)

def save_to_word(position, company, summary_text):
    """Save the summarized text to a Word document."""
    if not summary_text or summary_text.strip() == "":
        print("No content to save to document")
        return
        
    doc = Document()
    doc.add_heading(f"{position} Interview at {company}", level=1)
    doc.add_paragraph(f"Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    doc.add_paragraph("") 
    
    paragraphs = summary_text.split('\n')
    for para in paragraphs:
        if para.strip():
            doc.add_paragraph(para.strip())
    
    base_filename = f"{company}_{position.replace(' ', '_')}_Interview_{datetime.now().strftime('%Y%m%d_%H%M')}"
    sanitized_filename = sanitize_filename(base_filename) + ".docx"
    
    doc.save(sanitized_filename)
    print(f"Document saved: {sanitized_filename}")

def qa_pipeline(position, company):
    load_dotenv(dotenv_path=Path(__file__).with_name(".env"))
    API_KEY = os.getenv("SCRAPINGDOG_API_KEY")
    WATSONX_API_KEY = os.getenv("WATSONX_API_KEY")
    WATSONX_URL = os.getenv("WATSONX_URL")
    WATSONX_PROJECT_ID = os.getenv("WATSONX_PROJECT_ID")
    WATSONX_MODEL_ID = os.getenv("WATSONX_MODEL_ID")
    
    SCRAPINGDOG_URL = "https://api.scrapingdog.com/google/ai_mode"
    COUNTRY_CODE = get_location()
    
    try:
        search_results = search(position, company, COUNTRY_CODE, API_KEY, SCRAPINGDOG_URL)
        extracted_text = extract_text(search_results)
        
        if not extracted_text:
            print("No text content extracted from search results")
            exit(1)
        model = init_model(WATSONX_URL, WATSONX_API_KEY, WATSONX_PROJECT_ID, WATSONX_MODEL_ID)
        summary = summarize_content(model, extracted_text)
        
        if summary:
            save_to_word(position, company, summary)
        else:
            print("No summary generated")
            exit(1)
            
    except Exception as e:
        print(f"{e}")

if __name__ == "__main__":
    qa_pipeline()
